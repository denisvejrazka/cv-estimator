import pdfplumber
from docx import Document
import os

def extract_text(file_path: str) -> str:
    ext = os.path.splitext(file_path)[1].lower()
    
    if ext == ".pdf":
        return parse_pdf(file_path)
    elif ext == ".docx":
        return parse_docx(file_path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def parse_pdf(path: str) -> str:
    text = ""
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    
    if not text.strip():
        raise ValueError("PDF doesn't contain any text to extract")
    
    return text

# Extract from .docx (CV table aware approach)
def parse_docx(path: str) -> str:
    doc = Document(path)
    parts = [para.text for para in doc.paragraphs]
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                parts.append(cell.text)
    
    return "\n".join(filter(None, parts))

